"""Build the Groundwork beta testing guide as a real .docx."""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK   = RGBColor(0x11, 0x11, 0x11)
GREY  = RGBColor(0x66, 0x66, 0x66)
RED   = RGBColor(0xA0, 0x20, 0x20)

doc = Document()

# ── Page + base styles ───────────────────────────────────
s = doc.sections[0]
s.page_width, s.page_height = Cm(21.0), Cm(29.7)
for m in ('left_margin', 'right_margin'): setattr(s, m, Cm(2.4))
s.top_margin, s.bottom_margin = Cm(2.2), Cm(2.2)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

def shade(cell, hex_fill):
    el = OxmlElement('w:shd'); el.set(qn('w:fill'), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)

def fixed_layout(tbl):
    """Word ignores cell widths unless the table is explicitly fixed-layout."""
    tblPr = tbl._tbl.tblPr
    el = OxmlElement('w:tblLayout'); el.set(qn('w:type'), 'fixed')
    tblPr.append(el)

def borders(tbl, sz=4, color='BBBBBB'):
    tblPr = tbl._tbl.tblPr
    b = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz)); e.set(qn('w:color'), color)
        b.append(e)
    tblPr.append(b)

def rich(par, text, size=10.5, color=INK):
    """**bold** and `code` inside a paragraph."""
    for part in re.split(r'(\*\*.+?\*\*|`.+?`)', text):
        if not part: continue
        if part.startswith('**'):
            r = par.add_run(part[2:-2]); r.bold = True
        elif part.startswith('`'):
            r = par.add_run(part[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(size-1)
        else:
            r = par.add_run(part)
        r.font.size = Pt(size); r.font.color.rgb = color
    return par

def h1(text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = INK

def h2(text, page_break=True):
    p = doc.add_paragraph()
    if page_break: p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = INK
    pPr = p._p.get_or_add_pPr(); b = OxmlElement('w:pBdr'); bo = OxmlElement('w:bottom')
    bo.set(qn('w:val'),'single'); bo.set(qn('w:sz'),'8'); bo.set(qn('w:space'),'3'); bo.set(qn('w:color'),'111111')
    b.append(bo); pPr.append(b)

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = INK

def para(text, size=10.5, color=INK, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    return rich(p, text, size, color)

def meta(text):
    para(text, size=9, color=GREY, after=8)

def bullets(items, numbered=False):
    """Numbers are written as text, not as a Word list style.

    The built-in 'List Number' style shares one counter across the whole document, so
    Flow 1 began at 6 and every later flow drifted further. Restarting it properly means
    minting a numbering definition per list; writing '1.' is simpler and cannot drift.
    """
    for i, it in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.95)
        p.paragraph_format.first_line_indent = Cm(-0.55)
        marker = f'{i}.  ' if numbered else '\u2022  '
        rich(p, marker + it)

def callout(text, kind='info'):
    """One-cell table — survives Word far better than a bordered div."""
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    fixed_layout(t)
    t.columns[0].width = Cm(16.2)
    c = t.cell(0, 0); c.width = Cm(16.2)
    shade(c, 'FDF3F3' if kind == 'warn' else 'F4F4F4')
    borders(t, sz=6, color='A02020' if kind == 'warn' else 'BBBBBB')
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    rich(p, text, size=10, color=RED if kind == 'warn' else INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def no_split(row):
    """Keep a row on one page — a cell broken across a page break is unreadable."""
    el = OxmlElement('w:cantSplit')
    row._tr.get_or_add_trPr().append(el)

def repeat_header(row):
    """Repeat the header on every page a table spills onto.

    Without this, a table that crosses a page boundary leaves the continuation with
    unlabelled columns — 'Permit | Planning approval' with no idea what either column is.
    """
    el = OxmlElement('w:tblHeader'); el.set(qn('w:val'), 'true')
    row._tr.get_or_add_trPr().append(el)

def table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    fixed_layout(t)
    borders(t)
    no_split(t.rows[0]); repeat_header(t.rows[0])
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]; c.width = Cm(widths[i]); shade(c, 'EDEDED')
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
        r = p.add_run(htxt); r.bold = True; r.font.size = Pt(9.5)
    for row in rows:
        r_ = t.add_row(); no_split(r_)
        cells = r_.cells
        for i, val in enumerate(row):
            cells[i].width = Cm(widths[i])
            t.columns[i].width = Cm(widths[i])
            p = cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            rich(p, str(val), size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# =========================================================
h1('Groundwork by Jalla')
para('**Beta Testing Guide**', size=12, after=1)
meta('Version 1.0   ·   16 August 2026   ·   https://tryjalla.com')

callout('**What Groundwork is.**  A platform for people building a house in Africa from abroad. It estimates '
        'what the build should cost, breaks it into ten stages, holds each stage’s evidence and documents, '
        'and tracks payment against progress — so someone in London or Houston can tell whether the money '
        'they sent actually became a foundation.')

h2('1.  Before you start', page_break=False)

h3('What we need from you')
para('Use it as if it were your own money and your own house. Where something confuses you, that is the '
     'finding — please note what you expected instead. We are more interested in “I did not understand '
     'what this number meant” than in typos.')

h3('Accounts')
table(['Item', 'Detail'], [
    ['URL', 'https://tryjalla.com'],
    ['Sign up', 'Email + password, or Continue with Google'],
    ['Cost', 'Nothing. Stay on the free Self Verify plan throughout.'],
    ['Languages', 'English and French. Both are complete — please test both.'],
], [3.6, 12.6])

callout('**Please do not enter real payment details.**  Payment processing is not live. The Payments screens '
        'record what you say you have paid; they do not move money. Nothing in this beta will charge you.', 'warn')

h3('How to report something')
para('For each issue, please give us:')
bullets([
    'Which flow and step number from this document (e.g. “Flow 2, step 7”)',
    'What you expected to happen',
    'What actually happened',
    'A screenshot if the screen looked wrong',
    'Your browser, and whether you were on a phone or a computer',
], numbered=True)
para('If you see a red error message, please copy the exact wording. Error messages include a code in '
     'brackets — that code tells us precisely where it broke.')

# ── Flow 1 ──
h2('Flow 1 — Sign up and first sign-in')
meta('Expected time: 3 minutes')
bullets([
    'Go to `tryjalla.com`. You land on the marketing page.',
    'Click **Join for Free** (top right).',
    'Enter your name, email and a password. Or click **Continue with Google**.',
    'If you used email, check your inbox and click the confirmation link.',
    'You arrive at **Onboarding**. Tell it which country you are building in.',
    'You land on the **Dashboard**.',
], numbered=True)
h3('What to check')
bullets([
    'Does the confirmation email arrive within a minute? Does it land in spam?',
    'If you chose a French-speaking country (Cameroon, Senegal, Côte d’Ivoire), does the interface offer '
    'or switch to French? Is that helpful or annoying?',
    'Is the dashboard understandable when it is completely empty?',
    'Try **Log out** and sign back in.',
    'Try **Forgot password** — does the reset email arrive and work?',
])

# ── Flow 2 ──
h2('Flow 2 — Create a project  (the main flow)')
meta('Expected time: 10 minutes   ·   11 steps')
para('This is the most important flow in the beta. It is where we most need to know whether the questions '
     'make sense to someone who is not a builder.')
para('From the dashboard, click **New project**.')
table(['Step', 'Screen', 'What it asks'], [
    ['1',  'Country',         'Where you are building'],
    ['2',  'Project type',    'Residential, commercial, industrial, mixed use'],
    ['3',  'Building type',   'Single family, semi-detached, townhouse, multi-family…'],
    ['4',  'Floors',          'How many storeys'],
    ['5',  'Rooms',           'Bedrooms, bathrooms, living areas, kitchens, home offices — per floor'],
    ['6',  'Staff quarters',  'Yes or no'],
    ['7',  'Roof',            'First the form (pitched or flat), then the covering'],
    ['8',  'Details',         'City, floor area, finish level, project name, target start date'],
    ['9',  'Summary',         'Your build, costed trade by trade'],
    ['10', 'Plan',            'Self Verify / Jalla Verify / Jalla Management'],
    ['11', 'Confirm budget',  'Accept our estimate or type your contractor’s figure'],
], [1.3, 3.4, 11.5])

h3('Watch the drawing on the right')
para('As you answer, a technical drawing of your building assembles itself — floors stack, windows appear '
     'per room, the roof changes shape. **Does it match what you think you are describing?** If you say four '
     'bedrooms and see three windows, tell us.')

h3('Step 8 — the floor area is the one to scrutinise')
para('We suggest a floor area from your room list. This number drives the entire budget, so it matters more '
     'than anything else on the screen.')
callout('A 5-bedroom 2-storey semi-detached should suggest roughly **120 m² per floor**.\n\n'
        'Note this is the area of one floor, not the whole house. Is that clear from the screen? If you know '
        'what your house actually measures, does our suggestion look sensible?')

h3('Step 9 — the cost breakdown')
para('You will see the build priced trade by trade — foundation, structure, roof, joinery, electrical, '
     'plumbing, finishing — and then four totals:')
table(['Line', 'What it is'], [
    ['Construction',     'The build itself. Roughly 60% material, 40% labour.'],
    ['Design',           'Architectural and structural drawings, charged on built area'],
    ['Professional Fee', 'Groundwork’s fee: stage supervision and verification'],
    ['Permit',           'Planning approval and building permit'],
], [4.2, 12.0])
para('**Please check the four add up to the total shown.** They should, to the cent. And tell us whether you '
     'understand what you are paying for in each.')

h3('Step 11 — confirm the budget')
para('You can accept our estimate or type your own figure. Try typing a different number — say 10% higher. '
     'Every stage payment is recalculated from whatever you confirm.')

h3('What to check')
bullets([
    'Does any question assume knowledge you do not have?',
    'Use **Back** and change an earlier answer — does everything after it update?',
    'Leave the wizard half-finished, close the tab, come back. Are your answers still there?',
    'Do it on a phone. Is anything unusable?',
    'Switch to French mid-wizard. Does anything stay in English?',
])
callout('**Free plan limit: three active projects.**  To create a fourth, archive an old one first (Flow 8). '
        'Archived projects do not count.')

# ── Flow 3 ──
h2('Flow 3 — Track a project')
meta('Expected time: 10 minutes')
para('Open your project. Ten stages, in the order a building is actually built:')
table(['#', 'Stage', '#', 'Stage'], [
    ['1', 'Land Secured',      '6',  'Roofing'],
    ['2', 'Design Completed',  '7',  'Electrical & Plumbing'],
    ['3', 'Site Preparation',  '8',  'Finishing'],
    ['4', 'Foundation',        '9',  'Exterior Work'],
    ['5', 'Structure & Walls', '10', 'Final Handover'],
], [1.0, 7.1, 1.0, 7.1])
para('Each stage holds several substages — 60 in total. Only the current stage is unlocked; you cannot '
     'skip ahead, by design.')
bullets([
    'Open the **Stages** tab.',
    'Expand Stage 1. You will see its substages (survey, title, notary…).',
    'Upload a photo or document as evidence against a substage. Any image will do.',
    'Mark the substage complete.',
    'Complete every substage in the stage, then approve the stage.',
    'Stage 2 unlocks.',
], numbered=True)
h3('What to check')
bullets([
    'Is it obvious what evidence each substage wants?',
    'Can you tell at a glance which stage you are on and what is left?',
    'Try uploading a large photo straight from a phone camera. Does it work?',
    'Do the stage names match how a builder in your country would describe them?',
])
callout('**Land Secured and Exterior Work carry no money.**  Land is bought before Groundwork is involved, '
        'and exterior work is not in the budget. They are still tracked as stages. Is that clear, or does a '
        '$0 stage look like a bug?')

# ── Flow 4 ──
h2('Flow 4 — Budget and costing')
meta('Expected time: 5 minutes')
bullets([
    'Open the **Costing** tab.',
    'Click **“How is this calculated?”** — a step-by-step derivation of every figure.',
    'Expand the **Construction** line to see it spread across the ten stages.',
    'Click **Export PDF**.',
], numbered=True)
h3('What to check')
bullets([
    '**Do the numbers agree everywhere?** The total on the Overview donut, the Costing tab, the Payments '
    'schedule and the PDF must all be identical. If any two differ, that is a serious bug — please '
    'screenshot both.',
    'Does the explanation actually explain it, or does it restate the number?',
    'Is the PDF something you would send to a contractor or a bank?',
])

# ── Flow 5 ──
h2('Flow 5 — Payments')
meta('Expected time: 5 minutes')
para('Open the **Payments** tab. Each stage carries a payment milestone. Design, permit and professional fees '
     'appear as their own lines because they are not tied to site work.')
bullets([
    'Record a payment against Stage 1.',
    'Watch the paid total and the progress bar update.',
    'Mark it unpaid again.',
], numbered=True)
h3('What to check')
bullets([
    'Do the milestones add up to your confirmed budget?',
    'Are amounts shown in your local currency as well as dollars? Is the rate believable?',
    'Is it clear this is a record, not an actual payment?',
])

# ── Flow 6 ──
h2('Flow 6 — Documents and messages')
meta('Expected time: 5 minutes')
bullets([
    '**Documents** tab — upload a PDF and an image. Rename one. Download one. Delete one.',
    '**Messages** tab — post a message. This is where you would talk to your contractor.',
], numbered=True)
h3('What to check')
bullets([
    'Can you tell which documents belong to which stage?',
    'Does upload work from a phone?',
    'Is the 500 MB limit on the free plan visible before you hit it?',
])

# ── Flow 7 ──
h2('Flow 7 — Contractors')
meta('Expected time: 10 minutes   ·   needs two people or two accounts')
h3('7a — As the homeowner')
bullets([
    'Open your project, find **Invite contractor**.',
    'Enter an email address and send the invitation.',
], numbered=True)
para('The free plan allows one contractor per project.')
h3('7b — As the contractor')
bullets([
    'Open the invitation email and follow the link.',
    'Create an account or sign in.',
    'You see the project, but only the Stages and Messages tabs — not the owner’s budget.',
], numbered=True)
h3('7c — Bill of Quantities')
para('Contractors can price the job line by line. Open **Bill of Quantities** from the project.')
bullets([
    'Click **Start a take-off**. Every line arrives already priced.',
    'Optionally enter the building’s real length and width.',
    'Change a rate — say blockwork — to what you would actually pay. The total updates immediately.',
    'Use the reset arrow to put a line back to our figure.',
    '**Save draft**, then **Submit to client**.',
], numbered=True)
h3('What to check (contractors especially)')
bullets([
    '**Do our item numbers match how you write a BQ?** We use 204 footings, 305 blockwork, 503 roof sheet, '
    '801–810 plumbing.',
    'Lines marked “est.” are rates we inferred rather than measured. Are any badly wrong?',
    'Is anything missing that you would always price?',
    'Is working in XAF with dollars alongside the right way round?',
])
callout('A draft take-off is private to you. Only once submitted does the homeowner see it, and the figures '
        'are then frozen so later price changes cannot alter what you quoted.')

# ── Flow 8 ──
h2('Flow 8 — Archive and delete')
meta('Expected time: 3 minutes')
para('At the foot of any project page:')
table(['Action', 'What happens'], [
    ['Archive', 'Hides it and frees a slot on your plan. Everything is kept. Reversible.'],
    ['Delete',  'Permanent. Removes the project, its stages, documents, payments and messages.'],
], [3.2, 13.0])
bullets([
    'Archive a project. Check your project count drops and you can create a new one.',
    'Restore it. Check the count goes back up.',
    'Delete a test project you do not need. Read the warning before confirming.',
], numbered=True)
h3('What to check')
bullets([
    'Is the difference between archive and delete clear before you click?',
    'Does the delete warning tell you enough to decide?',
    'Did anything survive that should not have?',
])

# ── Flow 9 ──
h2('Flow 9 — Free tools  (no account needed)')
meta('Expected time: 5 minutes')
para('At `tryjalla.com/tools`, open to anyone:')
bullets([
    '**Build Budget Calculator** — a cost estimate from country, area, floors and finish',
    '**Stage Guide** — the ten stages explained',
    '**Milestone Planner** — how a budget splits across stages',
    '**Project Tracker** — a sample of the tracking view',
])
para('**Would these persuade you to sign up?** That is what we are really asking.')

# ── Flow 10 ──
h2('Flow 10 — Language')
meta('Expected time: 5 minutes   ·   French speakers especially')
para('Switch between EN and FR using the toggle in the header. Please walk one complete flow in French.')
h3('What to check')
bullets([
    'Any English text left behind on a French screen?',
    'Are the construction terms the ones actually used on site in Cameroon? We would rather have the word a '
    'mason uses than a dictionary translation.',
    'Do numbers and currency read correctly?',
    'Does any text overflow its button now that it is longer?',
])

# ── Known limits ──
h2('Known limits — please do not report these')
table(['Area', 'Status'], [
    ['Payments',        'Recording only. No money moves. Stripe and mobile money are not connected.'],
    ['Cost accuracy',   'Calibrated against four real Cameroonian bills of quantities. Cameroon is the most '
                        'accurate; other countries are indexed estimates.'],
    ['Staff quarters',  'We ask, but do not price it yet. Awaiting a bill of quantities.'],
    ['Some BQ rates',   'Lines marked “est.” have no source document behind them yet.'],
    ['Nigeria',         'Rates are unverified. We have no Nigerian bill of quantities.'],
    ['Jalla Management','Budget is set by our team after you create the project, so it opens showing a banner.'],
], [3.6, 12.6])

# ── Appendix ──
h2('Appendix — recording a walkthrough video')
para('If you are recording a demo rather than testing, this is the running order. Roughly 8 minutes at a '
     'comfortable pace. Record at 1920×1080, browser full screen, zoom at 100%.')
table(['#', 'Segment', 'Length', 'Show'], [
    ['1',  'Landing',           '0:30', 'tryjalla.com, scroll the page slowly'],
    ['2',  'Free tool',         '0:45', '/tools/budget — drag the area slider, watch the four lines move'],
    ['3',  'Sign up',           '0:45', 'Join for Free → form → confirmation → onboarding'],
    ['4',  'Empty dashboard',   '0:20', 'What a new user sees'],
    ['5',  'Wizard 1–7',   '2:00', 'Keep the drawing on the right in shot as it builds'],
    ['6',  'Wizard 8–9',   '1:00', 'Pause on the floor-area suggestion and the cost breakdown'],
    ['7',  'Wizard 10–11', '0:45', 'Pick Self Verify, confirm the budget, create'],
    ['8',  'Project overview',  '0:45', 'Donut, stage list, “How is this calculated?”'],
    ['9',  'Track a stage',     '1:00', 'Upload evidence, complete substages, approve, stage 2 unlocks'],
    ['10', 'Payments',          '0:30', 'Record a payment, watch the bar move'],
    ['11', 'Contractor BQ',     '0:45', 'Open a take-off, change one rate, total updates'],
    ['12', 'French',            '0:15', 'Toggle FR on the project page'],
], [1.0, 3.6, 1.8, 9.8])
callout('**Before recording:** use a throwaway account, not a real client’s. The recording will show the '
        'account email and any project names on screen. Archive or delete the demo project afterwards.', 'warn')

p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(18)
r = p.add_run('Groundwork by Jalla  ·  Beta Testing Guide  ·  v1.0  ·  16 August 2026')
r.font.size = Pt(8.5); r.font.color.rgb = GREY

doc.save('/home/favour-nwachukwu/Desktop/Jalla/groundwork1/docs/Groundwork-Beta-Testing-Guide.docx')
print('saved')
