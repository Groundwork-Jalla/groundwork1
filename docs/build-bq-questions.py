"""Build the BQ questions for Vanessa as a .docx.

Written for a quantity surveyor, not a developer: no software vocabulary, every question
anchored to a figure from her own documents, and a space to write the answer.
"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK, GREY, RED = RGBColor(0x11,0x11,0x11), RGBColor(0x66,0x66,0x66), RGBColor(0xA0,0x20,0x20)
doc = Document()

s = doc.sections[0]
s.page_width, s.page_height = Cm(21.0), Cm(29.7)
s.left_margin = s.right_margin = Cm(2.3)
s.top_margin = s.bottom_margin = Cm(2.0)

n = doc.styles['Normal']
n.font.name = 'Calibri'; n.font.size = Pt(10.5); n.font.color.rgb = INK
n.paragraph_format.space_after = Pt(6); n.paragraph_format.line_spacing = 1.15

def shade(cell, hexfill):
    el = OxmlElement('w:shd'); el.set(qn('w:fill'), hexfill)
    cell._tc.get_or_add_tcPr().append(el)

def fixed(tbl):
    el = OxmlElement('w:tblLayout'); el.set(qn('w:type'), 'fixed')
    tbl._tbl.tblPr.append(el)

def borders(tbl, sz=4, color='BBBBBB'):
    b = OxmlElement('w:tblBorders')
    for e in ('top','left','bottom','right','insideH','insideV'):
        x = OxmlElement(f'w:{e}')
        x.set(qn('w:val'),'single'); x.set(qn('w:sz'),str(sz)); x.set(qn('w:color'),color)
        b.append(x)
    tbl._tbl.tblPr.append(b)

def no_split(row):
    row._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))

def repeat(row):
    el = OxmlElement('w:tblHeader'); el.set(qn('w:val'),'true')
    row._tr.get_or_add_trPr().append(el)

def rich(par, text, size=10.5, color=INK, bold=False):
    for part in re.split(r'(\*\*.+?\*\*)', text):
        if not part: continue
        r = par.add_run(part[2:-2] if part.startswith('**') else part)
        r.bold = bold or part.startswith('**')
        r.font.size = Pt(size); r.font.color.rgb = color
    return par

def para(t, size=10.5, color=INK, after=6, bold=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    return rich(p, t, size, color, bold)

def h1(t):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(t); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = INK

def h2(t, brk=True):
    p = doc.add_paragraph()
    if brk: p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(7)
    r = p.add_run(t); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = INK
    pPr = p._p.get_or_add_pPr(); bd = OxmlElement('w:pBdr'); bo = OxmlElement('w:bottom')
    bo.set(qn('w:val'),'single'); bo.set(qn('w:sz'),'8'); bo.set(qn('w:space'),'3'); bo.set(qn('w:color'),'111111')
    bd.append(bo); pPr.append(bd)

def h3(t):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(t); r.bold = True; r.font.size = Pt(12.5); r.font.color.rgb = INK

def bullets(items, numbered=False, size=10.5):
    for i, it in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.95)
        p.paragraph_format.first_line_indent = Cm(-0.55)
        rich(p, (f'{i}.  ' if numbered else '•  ') + it, size)

def table(headers, rows, widths, size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.autofit = False
    fixed(t); borders(t); no_split(t.rows[0]); repeat(t.rows[0])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.width = Cm(widths[i]); t.columns[i].width = Cm(widths[i])
        shade(c, 'EDEDED')
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
        r = p.add_run(h); r.bold = True; r.font.size = Pt(size)
    for row in rows:
        rr = t.add_row(); no_split(rr)
        for i, v in enumerate(row):
            rr.cells[i].width = Cm(widths[i])
            p = rr.cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            rich(p, str(v), size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def box(text, kind='info'):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.autofit = False; fixed(t)
    t.columns[0].width = Cm(16.4); c = t.cell(0,0); c.width = Cm(16.4)
    shade(c, 'FDF3F3' if kind=='warn' else 'F4F4F4')
    borders(t, 6, 'A02020' if kind=='warn' else 'BBBBBB')
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    rich(p, text, 10, RED if kind=='warn' else INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def answer_box(height_lines=3):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.autofit = False; fixed(t)
    t.columns[0].width = Cm(16.4); c = t.cell(0,0); c.width = Cm(16.4)
    borders(t, 4, '999999')
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    r = p.add_run('Your answer'); r.font.size = Pt(8.5); r.font.color.rgb = GREY
    for _ in range(height_lines):
        c.add_paragraph().paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def question(num, title, sees, assumed, need, answer_lines=3, evidence=None):
    h3(f'Question {num}  —  {title}')
    if sees:
        para('**What your document shows**', size=10, after=2)
        for s_ in sees: para(s_, size=10, after=3)
    if evidence:
        table(*evidence)
    if assumed:
        para('**What we have assumed in the meantime**', size=10, after=2)
        para(assumed, size=10, after=4)
    para('**What we need from you**', size=10, after=2)
    rich(doc.add_paragraph(), need, 10.5, INK, bold=True)
    answer_box(answer_lines)

# ═════════════════════════════════════════════════════════
h1('Bill of Quantities — questions')
para('**For Vanessa Gwanvoma**', size=11.5, after=1)
para('Groundwork by Jalla   ·   17 August 2026', size=9.5, color=GREY, after=8)

box('**What this is.**  We built Groundwork’s cost estimator from four of your bills of '
    'quantities. It now reproduces them to within about 21% under and 31% over, against '
    '146% over when it was built from a single document.\n\n'
    'The remaining gap sits in a small number of specific line items where your four '
    'documents disagree with one another, and we cannot tell from the files alone which '
    'reading is the correct one. Each question below is one of those. Every one you '
    'answer lets us tighten a real figure.')

h3('How to answer')
para('There is a box under each question. A sentence is plenty — we are not asking you to '
     'reprice anything. Where a question offers two readings, telling us which one is right '
     'is the whole answer. Where we have simply guessed a rate, a corrected figure is all we '
     'need.')
para('If it is easier to talk than write, we will take notes on a call and send them back '
     'for you to correct.')

h3('The four documents')
table(['#','File','Project','Total (XAF)'], [
    ['1','CONSTRUCTION ESTIMATE OF G+1 BUILDING .xlsx','Mme. Rose Ndum Kenah, Yaoundé','59,675,280'],
    ['2','BUEA RESIDENCE ESTIMATE.xlsx','Woyamukumbat, Buea','43,410,955'],
    ['3','NI PAS CONSTRUCTION ESTIMATE OF G+1 SCHOOL.xlsx','Naka','42,213,867'],
    ['4','MPANGOU.xlsx','Mpangou, Kribi (G+3)','64,268,593'],
], [0.9, 7.2, 4.6, 3.7])

h3('Where our estimate currently lands against each')
para('These are **construction against construction**. Your bills price the build itself, so '
     'the figures below are our build cost only — they exclude the design, permit and '
     'professional fees Groundwork charges a client on top. Comparing our client-facing '
     'total against your documents would flatter or penalise us for fees that are not in '
     'your scope.', size=10, after=6)
table(['Document','Yours (XAF)','Ours (XAF)','Difference'], [
    ['3  Naka',    '42,213,867','39,943,940','**−5.4%**'],
    ['2  Buea',    '43,410,955','51,027,448','+17.5%'],
    ['1  Rose',    '59,675,280','47,301,063','−20.7%'],
    ['4  Mpangou', '64,268,593','84,228,987','+31.1%'],
], [4.2, 4.2, 4.2, 3.8])
para('Naka is closest because it is one of the two documents that measured internal '
     'partitions and painted the whole building. The two widest, Rose and Mpangou, are the '
     'two carrying the anomalies in Questions 1, 3 and 4 — which is why those three matter '
     'most.', size=10, color=GREY)


# ── Briefing ─────────────────────────────────────────────
def review(num, title, body, need, answer_lines=4):
    h3(f'Review {num}  —  {title}')
    for b in body: para(b, size=10, after=3)
    para('**What we need from you**', size=10, after=2)
    rich(doc.add_paragraph(), need, 10.5, INK, bold=True)
    answer_box(answer_lines)

h2('How we arrive at a figure')
para('This section is here so you can check the shape of the model, not only the individual '
     'line items. If the way we spread money across a build does not match how it is really '
     'spent on site, everything else is detail.')

h3('What a client is quoted')
para('Four separate things, shown as four lines. Only the first is a measured cost; the other '
     'three are fees.')
table(['Line','How it is worked out','On the example below'], [
    ['**Construction**','Measured quantities priced at your rates. The build itself.','$62,490.00'],
    ['**Design**','5,000 XAF per m² of built area (floor area × number of floors)','$2,000.00'],
    ['**Professional**','50,000 XAF per construction stage, seven stages — a flat fee','$583.33'],
    ['**Permit**','1% of the construction cost','$624.90'],
    ['**Total**','The four added together','**$65,698.23**'],
], [3.0, 9.4, 4.0])

h3('Where each of those numbers comes from')
para('This is the part worth your attention. **Only the first line is derived from your '
     'documents.** The rest are commercial decisions Jalla has taken, and we would like to '
     'know whether they look sane to someone who prices real buildings.')
table(['Figure','Where it came from','Your view?'], [
    ['The build cost','Your four bills of quantities — quantities from the geometry, rates as quoted in your documents','Yes — Questions 1–15'],
    ['60% material / 40% labour','A Jalla rule of thumb. Splits the build cost for display only; it does not change the total','Yes'],
    ['Design fee, 5,000 XAF/m²','Set by Jalla','Sanity check'],
    ['Professional fee, 50,000 XAF per stage','Set by Jalla','Sanity check'],
    ['Permit fee, 1% of the build','Set by Jalla','**Yes — Review 2**'],
    ['The ten stage percentages','Set by Jalla','**Yes — Review 1**'],
], [4.4, 8.0, 4.0])

h3('How the build cost itself is worked out')
para('In plain terms: we take the size and shape of the building, work out the quantities the '
     'way you would — cubic metres of excavation, square metres of blockwork, linear metres of '
     'railing — and price each at the rate your documents quote for that city.')
para('We do not apply a single rate per square metre to the whole house. That was the old '
     'method, fitted to one document, and it overshot the other three by up to 146%. '
     'Quantities first and rates second is what brought it into the range on the previous page.')

h3('How the money is spread across the ten stages')
para('A client pays in stages as work completes. Each stage carries a share of the **build '
     'cost** — not of the total, because the design, permit and professional fees are not site '
     'work and are paid separately.')
table(['#','Stage','Share of the build cost'], [
    ['1','Land Secured','—  never in the budget'],
    ['2','Design Completed','—  paid as the design fee'],
    ['3','Site Preparation','2%'],
    ['4','Foundation','8%'],
    ['5','Structure & Walls','30%'],
    ['6','Roofing','8%'],
    ['7','Electrical & Plumbing','17%'],
    ['8','Finishing','30%'],
    ['9','Exterior Work','—  not quoted'],
    ['10','Final Handover','5%'],
    ['','**The seven charged stages**','**100%**'],
], [1.2, 8.2, 7.0])
para('Three stages carry no percentage, for three different reasons. **Land** is bought before '
     'Groundwork is involved. **Design** is paid as its own fee at that stage rather than as a '
     'share. **Exterior work** sits outside what we quote — the budget covers the main building.')
para('The seven that remain add to exactly 100%, so the stage payments add up to the build cost '
     'with nothing unallocated and nothing counted twice.')

h3('A worked example')
para('A two-storey single-family house in Douala: 120 m² per floor, standard finish, four '
     'bedrooms, three bathrooms. These are real figures out of the system, not an illustration.')
table(['#','Stage','Share','Client pays'], [
    ['1','Land Secured','—','$0.00'],
    ['2','Design Completed','fee','$2,000.00'],
    ['3','Site Preparation','2%','$1,249.80'],
    ['4','Foundation','8%','$4,999.20'],
    ['5','Structure & Walls','30%','$18,747.00'],
    ['6','Roofing','8%','$4,999.20'],
    ['7','Electrical & Plumbing','17%','$10,623.30'],
    ['8','Finishing','30%','$18,747.00'],
    ['9','Exterior Work','—','$0.00'],
    ['10','Final Handover','5%','$3,124.50'],
    ['','Permit fee','1%','$624.90'],
    ['','Professional fee','flat','$583.33'],
    ['','**Total the client pays**','','**$65,698.23**'],
], [1.2, 7.2, 2.4, 5.6])
para('The seven charged stages add to $62,490.00 — the build cost exactly. Add the design fee '
     'at stage 2, then the permit and professional fees as their own lines, and you have the '
     '$65,698.23 the client sees.', size=10, color=GREY)

review(1, 'Do those stage percentages match how money is really spent?',
    ['This weighting came from Jalla, not from your documents. It decides how much a client '
     'hands over at each point in the build.',
     'Getting it wrong is not a rounding problem. Too little early and a contractor cannot buy '
     'materials to start; too much early and the client has paid for work that has not '
     'happened, which is the exact risk Groundwork exists to remove.',
     'The seven charged stages are: site preparation 2, foundation 8, structure and walls 30, '
     'roofing 8, electrical and plumbing 17, finishing 30, final handover 5.'],
    'Does that spread look right for a Cameroonian residential build? If not, what would you '
    'change it to — the seven need to total 100.')

review(2, 'Is 1% of the build a fair figure for the permit?',
    ['We charge the client 1% of the construction cost to cover planning approval, the building '
     'permit and lands registry.',
     'On the worked example that is $624.90, or roughly 375,000 XAF.'],
    'Is that about right in Cameroon, and does it vary much by city or by the size of the '
    'building?')

# ── Section A ────────────────────────────────────────────
h2('Section A  —  Figures in the documents we could not read')
para('These five are places where a quantity in one document sits far outside the range of '
     'the other three. In each case we suspect a unit or a scope difference rather than an '
     'error, but we cannot tell which.')

question(1, 'Roof timber on the Rose document',
    ['Item 501 lists **806.20 m³** of treated timber, and item 502 a further **520 m³**, '
     'on a 125 m² house. Taken literally that is roughly 500 tonnes of wood.',
     'The same items in your other three documents are between 57 and 90 m³.',
     'Roof cost works out at **48,050 XAF per m² of footprint** here, against 5,202 / 6,246 / '
     '9,804 in the others.'],
    'We treat this figure as an outlier and price a pitched roof at about 20,000 XAF per m² '
    'of footprint from first principles. That is a large part of why our estimate reads 20.7% '
    'below your Rose total.',
    'Is 806.20 linear metres, or board feet, recorded in a column headed m³?')

question(2, 'Roof sheeting on the Buea document',
    ['Item 503 lists **42.67 m²** of aluminium sheet on a **224 m²** footprint. A roof over '
     'that footprint would normally need 250–290 m² of sheet.',
     'The document also carries a parapet line (336 m³ at 780 XAF).'],
    'We have read this as a flat concrete roof where only a small canopy is sheeted, because '
    'of the parapet line.',
    'Is that right, or is a roofing section missing from the file?')

question(3, 'How many floors is Mpangou?',
    ['The section is headed **“FIRST TILL 3RD FLOOR ELEVATION”** and the file is titled G+3.',
     'But the section total of 24,514,368 is exactly **four times** one floor’s line items '
     '(6,128,592). In your other three documents the section total is exactly one times its '
     'line items.'],
    'We have taken it as stated in the title. If it is actually four upper floors, the '
    'per-floor uplift we derive changes from 20.6% to 15.4%, and this is a direct cause of '
    'our +31.1% on this document.',
    'Is Mpangou three upper floors, or four?')

question(4, 'The Mpangou foundation',
    ['Against document 3, on the same 144 m² footprint and half the height:'],
    'We have excluded this document from our foundation calibration entirely, because a '
    'four-storey building needs larger footings than a two-storey one on the same footprint, '
    'not sixteen times smaller.',
    'Was this foundation designed for a different structure, or is the schedule incomplete?',
    evidence=(['Item','Mpangou (G+3)','Naka (G+1)','Ratio'], [
        ['Excavation','40.32 m³','112.32 m³','0.36×'],
        ['Reinforced concrete footings (204)','0.69 m³','11.66 m³','**0.06×**'],
        ['Foundation beams (206)','2.16 m³','6.48 m³','0.33×'],
        ['Preliminaries (100)','170,000','830,000','0.20×'],
    ], [6.0, 3.6, 3.6, 3.2]))

question(7, 'Painting scope on Mpangou',
    ['Items 901 and 902 cover **259.20 m²**, which is one floor’s plastered area, on a '
     'four-storey building.',
     'Your other three documents paint every floor.'],
    'We paint every floor in our own calculation.',
    'Should Mpangou’s painting be four times 259.20 m²?')

# ── Section B ────────────────────────────────────────────
h2('Section B  —  Conventions we had to choose between')
para('These four are not errors in any document. They are places where your documents follow '
     'different conventions, and we had to pick one. Knowing which is standard practice would '
     'let us stop guessing.')

question(5, 'What unit is blockwork priced in?',
    ['Document 1 prices item 305 at **7,800 XAF per m²**, with quantities in m² '
     '(412.80 ground, 370 first), and its plastering is exactly twice that — both faces of '
     'the wall.',
     'Documents 2, 3 and 4 price the same item at **1,750 XAF**, against quantities of '
     '1,456 / 936 / 312.'],
    'We have adopted the document 1 convention: per m² of wall. Neither reading reconciles '
    'with the plastered areas in the other three documents.',
    'Is 1,750 XAF per block, or per m²? This now matters more than it did — contractors '
    'using Groundwork see a unit printed next to the box they are editing, and they need to '
    'know what they are pricing.')

question(6, 'Do documents 2 and 4 exclude internal partitions?',
    ['Plastered area per floor, against the external envelope counted on both faces:'],
    'We model internal partitions at 14 metres of wall per room per floor. Documents 1 and 3 '
    'imply 15.5 m and 18 m. That is most of why our estimate reads 17.5% above your Buea '
    'total.',
    'Do documents 2 and 4 deliberately exclude internal partition walls, or were they '
    'omitted?',
    evidence=(['Document','Plastered per floor','External envelope, both faces','Partitions included?'], [
        ['1  Rose','825.60 m²','268 m²','yes'],
        ['3  Naka','777.60 m²','288 m²','yes'],
        ['2  Buea','388.80 m²','360 m²','**no**'],
        ['4  Mpangou','259.20 m²','288 m²','**no**'],
    ], [3.4, 4.2, 5.2, 3.6]))

question(8, 'Which specification counts as “standard”?',
    ['Normalised for city and floor area, shell cost per built m² runs from 112 to 232 USD '
     'across your four documents.',
     'Document 1 has ceiling staffing on both floors, two bathtubs, five mirrors and 1.5M of '
     'decoration. Documents 2 and 3 have none of those.'],
    'We classify documents 2 and 3 as standard finish and document 1 as premium, on that '
    'basis. This sets the multipliers a client sees when they choose a finish level: '
    '**standard 1.00, premium 1.45, luxury 1.70**.',
    'Is that the right reading, and are those three multipliers about right?')

question(9, 'Is electrical a quoted package or a measured install?',
    ['Section 700 totals **2,326,600 XAF in documents 1, 2 and 3 — identical to the franc** '
     '— despite those buildings having 12, 9 and 9 rooms.',
     'Section 800 (plumbing) is likewise identical in documents 2 and 3.'],
    'We scale electrical by the number of floors, with a small addition per room, so that a '
    'ten-bedroom house still prices above a two-bedroom one.',
    'Are these standard packages you apply, or is the match a coincidence? If they are fixed '
    'packages we should label those lines “package — quoted” and show no quantity, rather '
    'than showing a client a measurement we did not measure.')

# ── Section C ────────────────────────────────────────────
h2('Section C  —  New questions since we last wrote')
para('Groundwork now prints its estimate as a line-by-line bill using your item numbering, so '
     'a contractor can lay it beside a real quotation. That change surfaced these six.')

question(10, 'Mpangou lists no mirrors',
    ['Our fixture schedule reproduces documents 1, 2 and 3 **to the franc** '
     '(6,450,000 / 2,255,000 / 2,255,000).',
     'Document 4 is the one it misses. It is plainly a luxury specification — five bathtubs, '
     'eight showers, five kitchen sinks — but lists **no mirrors at all** under item 807.'],
    'We give a mirror to every bathroom on a premium or luxury specification, so we show three '
    'where your document shows none. We have deliberately not changed the rule to fit, because '
    'we do not know which is correct.',
    'Is that an omission in the document, or do you not price mirrors on this kind of job?')

question(11, 'Is Adamawa 7–8% overall, or 7–8% on labour and materials?',
    ['You told us Adamawa runs 7–8% above the Douala baseline, and we have corrected our '
     'index from 1.4444 to **1.0750** accordingly.',
     'But an identical building still prices **15.2%** above Douala, because we take concrete '
     'from each city’s own column and Adamawa’s is **260,000 XAF/m³ against Douala’s '
     '180,000** — 44% higher on inland haulage. Concrete is about a fifth of a take-off.'],
    'Both figures as you gave them: the index at 1.075 and the concrete column unchanged.',
    'Does your 7–8% describe the finished project, or only the non-concrete trades? If it '
    'is the finished project, the Adamawa concrete column is overstated too and the two need '
    'to move together.')

question(12, 'Roof coverings, and one we have added',
    ['We had three different sets of roof uplifts in our own system. They are now one set:'],
    'The first four are as we understood them from the documents. The fifth — a flat '
    'aluminium deck — appears in none of your four, and we have priced it at the long-span '
    'rate because it is the same sheet laid to a shallow fall rather than pitched.',
    'Are the first four right? And is our reading of the aluminium deck reasonable, or does '
    'the decking, falls and upstand detail put it nearer a concrete slab?',
    evidence=(['Covering','Uplift over long-span aluminium'], [
        ['Long-span aluminium','baseline'],
        ['Clay tiles','+10%'],
        ['Concrete slab','+8%'],
        ['Shingle','+5%'],
        ['**Aluminium deck (flat)**','**baseline — assumed, not measured**'],
    ], [7.0, 9.4]))

question(13, 'Room sizes behind your 120 m²',
    ['You put a five-bedroom two-storey semi-detached at **120 m² per floor**. We had been '
     'suggesting 95, so we rebuilt our room benchmarks to reach your figure:'],
    'These now drive every floor area the app suggests. A client who accepts our suggestion '
    'is accepting these numbers, so we would rather they were yours than ours.',
    'Are the individual room sizes right, or only the total? And is 12 m² sensible for a '
    'home office — that one is purely our guess.',
    evidence=(['Room','m²','Room','m²'], [
        ['Bedroom','17','Kitchen','15'],
        ['Bathroom','6','Home office','12  (our guess)'],
        ['Living room','32','Circulation','+30%'],
    ], [4.1, 4.1, 4.1, 4.1]))

question(14, 'Can you price a staff quarters?',
    ['We were adding **$8,000 per room** for staff quarters. That figure had nothing behind '
     'it, and on a typical build it was 19–27% of the total a client was shown.'],
    'We have removed it completely. The app still asks whether the client wants staff '
    'quarters, but it now costs nothing and says so on screen.',
    'Could you send a bill of quantities for a typical staff quarters block — say two rooms '
    'with a shared bathroom? Until then we cannot quote it at all.')

question(15, 'Nine rates that have nothing behind them',
    ['Now that our estimate prints as a line-by-line bill, every rate is visibly either taken '
     'from your documents or inferred by us. These nine are inferred, and they are flagged '
     '“est.” on screen so nobody mistakes them for measured figures:'],
    None,
    'Do any of these have a real rate you can give us? We would rather delete a line than '
    'publish an invented number inside it — a contractor who spots one made-up figure stops '
    'trusting the whole document.',
    answer_lines=4,
    evidence=(['Item','Description','What we currently assume'], [
    ['101','Site setup, hoarding, temporary works','250,000 XAF lump sum'],
    ['102','Site clearing and setting out','1,600 XAF/m²'],
    ['207','Damp-proof membrane','1,000 XAF/m²'],
    ['208','Sand blinding','1,000 XAF/m²'],
    ['501','Parapet wall','4,500 XAF/ml'],
    ['502','Roof timber and trusses','6,500 XAF/m²'],
    ['701','Electrical installation, per floor','700,000 XAF'],
    ['702','Electrical points, per room','90,000 XAF'],
    ['907','Final finishes allowance','500,000 XAF lump sum'],
    ], [1.8, 8.2, 6.4]))

# ── Price book ───────────────────────────────────────────
h2('The price book, and Nigeria')
para('The **unit cost calculation** sheet — eight cities, built up from cement, sand, gravel, '
     'steel, formwork, labour, water and equipment plus 25% overhead and profit — is now the '
     'basis of everything we price. It cross-checks well: each document’s own concrete rate '
     'matches its own city’s column.')
table(['Document','City','Concrete rate in the document','Matches the book?'], [
    ['1  Rose','Yaoundé','180,000','yes'],
    ['2  Buea','Buea','190,000','yes'],
    ['3  Naka','Bamenda','190,000','yes'],
    ['4  Mpangou','Kribi','179,000','yes'],
], [3.4, 3.2, 6.0, 3.8])
para('Two notes on that book. **Bali is now listed as Bamenda** in the app, carrying the same '
     'rates. And **Adamawa** is covered in Question 11 above.')

box('**Nigeria — we are guessing, and would rather not.**  Abuja sits at 450,000 XAF/m³ in '
    'the book, 2.5 times Douala. We currently hold three different Nigerian base rates in our '
    'own system — 672, 1,600 and 180 USD per m² — and not one of them has a Nigerian bill '
    'of quantities behind it.\n\n'
    '**Do you have a Nigerian BQ, or know someone who would share one?**  Until then we will '
    'keep saying on screen that Nigerian figures are unverified.', 'warn')
answer_box(3)

# ── What answers unlock ──────────────────────────────────
h2('What each answer would let us do')
table(['Questions','What changes'], [
    ['1, 3, 4, 6, 7','We stop excluding documents from calibration, and can tighten our stated accuracy from roughly ±25–35% toward ±10%'],
    ['5, 9, 15','Contractors can price against our lines without translating them into their own first'],
    ['10','Our fixture schedule reproduces all four of your documents instead of three'],
    ['11','Clients building in Adamawa get a figure we can defend'],
    ['12, 13','Every roof and floor area the app suggests rests on your figures rather than ours'],
    ['14','We can quote staff quarters at all'],
], [3.6, 12.8])

para('If the re-baselined bill of quantities you mentioned is close to ready, several of these '
     'would be settled by it at once — we are happy to wait for that rather than take your '
     'time twice.', after=10)
para('Thank you. Every question here exists because we would rather ask than guess.',
     size=10.5, bold=True)
para('Groundwork by Jalla   ·   contact@tryjalla.com', size=9, color=GREY)

doc.save('/home/favour-nwachukwu/Desktop/Jalla/groundwork1/docs/Groundwork-BQ-Questions-Vanessa.docx')
print('saved')
